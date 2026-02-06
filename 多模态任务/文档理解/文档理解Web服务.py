#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档理解 Web 服务 - 文档解析师 📄
"""

import os
os.environ['HF_HOME'] = r'D:\transformers训练\transformers-main\预训练模型下载处'
os.environ['TRANSFORMERS_CACHE'] = r'D:\transformers训练\transformers-main\预训练模型下载处'
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'  # 使用国内镜像
# 不设置TESSDATA_PREFIX，因为路径包含中文会导致Tesseract无法识别

from flask import Flask, request, jsonify, render_template_string, send_file
from transformers import pipeline
from PIL import Image
import base64
import io
import tempfile

# 配置 Tesseract-OCR 并修复编码问题
try:
    import pytesseract
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    
    # Monkey patch pytesseract 修复Windows中文编码问题
    import pytesseract.pytesseract as pyt
    
    # 修改get_errors函数，使其能处理GBK编码
    original_get_errors = pyt.get_errors
    
    def patched_get_errors(error_bytes):
        """修复编码错误 - 尝试多种编码"""
        if isinstance(error_bytes, bytes):
            # 尝试多种编码
            for encoding in ['utf-8', 'gbk', 'gb2312', 'cp936', 'latin1']:
                try:
                    return error_bytes.decode(encoding)
                except (UnicodeDecodeError, AttributeError):
                    continue
            # 所有编码都失败，使用ignore模式
            return error_bytes.decode('utf-8', errors='ignore')
        return error_bytes
    
    pyt.get_errors = patched_get_errors
    
    # 同时修改run_and_get_output，默认使用中英文双语言
    original_run_and_get_output = pyt.run_and_get_output
    
    def patched_run_and_get_output(image, extension, lang=None, config='', nice=0, timeout=0, return_bytes=False):
        """默认使用中英文双语言（语言包已安装到Tesseract默认目录）"""
        # 如果没有指定语言或只指定了英文，使用中英文双语言
        if lang is None or lang == 'eng':
            lang = 'chi_sim+eng'  # 中文+英文
        
        try:
            return original_run_and_get_output(image, extension, lang, config, nice, timeout, return_bytes)
        except Exception as e:
            # 如果中文语言包失败，降级到只用英文
            if 'chi_sim' in str(e) and lang == 'chi_sim+eng':
                print(f"⚠️ 中文语言包加载失败，降级到只用英文: {e}")
                lang = 'eng'
                return original_run_and_get_output(image, extension, lang, config, nice, timeout, return_bytes)
            raise
    
    pyt.run_and_get_output = patched_run_and_get_output
    
    print("✅ pytesseract编码补丁已应用（支持中英文双语言）")
except ImportError:
    print("⚠️ pytesseract未安装")
except Exception as e:
    print(f"⚠️ pytesseract补丁应用失败: {e}")

# 导入PDF和Word处理库
try:
    from pdf2image import convert_from_path
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    print("⚠️  未安装 pdf2image，PDF支持不可用。安装: pip install pdf2image")

try:
    from docx2pdf import convert as docx_to_pdf
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    print("⚠️  未安装 docx2pdf，Word支持不可用。安装: pip install docx2pdf")

try:
    from docx import Document
    DOCX_READ_SUPPORT = True
    print("✅ python-docx 已安装，Word读取支持已启用")
except ImportError:
    DOCX_READ_SUPPORT = False
    print("⚠️  未安装 python-docx，Word读取支持不可用。")
    print("   安装命令: pip install python-docx")

try:
    import pythoncom
    PYTHONCOM_SUPPORT = True
except ImportError:
    PYTHONCOM_SUPPORT = False

app = Flask(__name__)

CURRENT_DIR = os.path.dirname(os.path.abspath(__file__))
BACKGROUND_PATH = os.path.join(CURRENT_DIR, '背景.png')

print("=" * 70)
print("📄 文档理解 Web 服务 - 文档解析师")
print("=" * 70)

# 检查 Tesseract-OCR
print("\n🔍 检查 Tesseract-OCR...")
try:
    import pytesseract
    # 配置 Tesseract 路径
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    
    # 检查语言数据文件
    tessdata_path = r'D:\transformers训练\transformers-main\预训练模型下载处'
    eng_data = os.path.join(tessdata_path, 'eng.traineddata')
    
    if not os.path.exists(eng_data):
        print("❌ Tesseract语言数据文件缺失！")
        print(f"   需要下载 eng.traineddata 到 {tessdata_path}")
        TESSERACT_AVAILABLE = False
    else:
        print(f"✅ 语言数据文件已找到: {eng_data}")
        # 尝试运行 tesseract
        pytesseract.get_tesseract_version()
        print("✅ Tesseract-OCR 已安装并配置")
        TESSERACT_AVAILABLE = True
except Exception as e:
    print(f"⚠️  Tesseract-OCR 配置失败: {e}")
    TESSERACT_AVAILABLE = False

print("\n📚 正在加载文档问答模型...")
doc_qa = pipeline("document-question-answering", model="impira/layoutlm-document-qa")
print("✅ 文档解析师准备完毕！")

# 打印支持的格式
print("\n📋 支持的文档格式:")
print("  ✅ 图片格式: JPG, PNG, BMP, GIF")
if PDF_SUPPORT:
    print("  ✅ PDF文档")
else:
    print("  ❌ PDF文档 (需要安装 pdf2image 和 poppler)")
if DOCX_READ_SUPPORT:
    print("  ✅ Word文档 (.docx)")
else:
    print("  ❌ Word文档 (需要安装 python-docx)")

HTML_TEMPLATE = """
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>📄 文档理解 - 文档解析师</title>
    <style>
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }
        
        body {
            font-family: 'Microsoft YaHei', 'Arial', sans-serif;
            background-image: url('/static/background');
            background-size: cover;
            background-position: center;
            background-attachment: fixed;
            min-height: 100vh;
            padding: 20px;
            overflow-y: auto;
            overflow-x: hidden;
        }
        
        .falling-item {
            position: fixed;
            font-size: 26px;
            animation: fall linear infinite;
            z-index: 1;
            pointer-events: none;
            opacity: 0.7;
        }
        
        @keyframes fall {
            0% {
                transform: translateY(-10px) rotate(0deg) scale(1);
                opacity: 0.7;
            }
            100% {
                transform: translateY(100vh) rotate(360deg) scale(1.26);
                opacity: 0.24;
            }
        }
        
        .container {
            background: linear-gradient(135deg, rgba(63, 81, 181, 0.95) 0%, rgba(48, 63, 159, 0.95) 100%);
            border-radius: 30px;
            box-shadow: 0 20px 60px rgba(63, 81, 181, 0.5);
            padding: 40px;
            max-width: 1100px;
            margin: 20px auto;
            width: 100%;
            backdrop-filter: blur(10px);
            border: 3px solid rgba(63, 81, 181, 0.6);
            position: relative;
            z-index: 10;
        }
        
        h1 {
            text-align: center;
            color: #fff;
            margin-bottom: 10px;
            font-size: 2.5em;
            text-shadow: 2px 2px 4px rgba(0,0,0,0.3);
        }
        
        .subtitle {
            text-align: center;
            color: #c5cae9;
            margin-bottom: 30px;
            font-size: 1.2em;
        }
        
        .upload-section {
            background: rgba(255, 255, 255, 0.95);
            border-radius: 20px;
            padding: 30px;
            margin-bottom: 25px;
        }
        
        .upload-area {
            border: 3px dashed #3f51b5;
            border-radius: 15px;
            padding: 40px;
            text-align: center;
            background: rgba(63, 81, 181, 0.05);
            cursor: pointer;
            transition: all 0.3s;
            margin-bottom: 20px;
        }
        
        .upload-area:hover {
            background: rgba(63, 81, 181, 0.1);
            border-color: #303f9f;
        }
        
        .upload-icon {
            font-size: 3.5em;
            margin-bottom: 15px;
        }
        
        #fileInput {
            display: none;
        }
        
        .preview-container {
            display: none;
            margin-bottom: 20px;
        }
        
        .preview-image {
            max-width: 100%;
            max-height: 400px;
            border-radius: 10px;
            box-shadow: 0 5px 15px rgba(0,0,0,0.2);
            display: block;
            margin: 0 auto 20px;
        }
        
        .question-area {
            margin-top: 20px;
        }
        
        .question-area label {
            display: block;
            color: #303f9f;
            font-weight: bold;
            margin-bottom: 10px;
            font-size: 1.1em;
        }
        
        input[type="text"] {
            width: 100%;
            padding: 15px;
            border: 2px solid #3f51b5;
            border-radius: 15px;
            font-size: 1.05em;
            font-family: 'Microsoft YaHei', 'Arial', sans-serif;
            transition: all 0.3s;
        }
        
        input[type="text"]:focus {
            outline: none;
            border-color: #303f9f;
            box-shadow: 0 0 15px rgba(63, 81, 181, 0.3);
        }
        
        .quick-questions {
            display: flex;
            gap: 10px;
            margin-top: 10px;
            flex-wrap: wrap;
        }
        
        .quick-btn {
            padding: 8px 15px;
            background: linear-gradient(135deg, #5c6bc0 0%, #3f51b5 100%);
            color: white;
            border: none;
            border-radius: 20px;
            cursor: pointer;
            font-size: 0.9em;
            transition: all 0.3s;
        }
        
        .quick-btn:hover {
            transform: translateY(-2px);
            box-shadow: 0 4px 12px rgba(63, 81, 181, 0.4);
        }
        
        .button-group {
            display: flex;
            gap: 15px;
            margin-top: 20px;
        }
        
        button {
            flex: 1;
            padding: 15px;
            font-size: 1.2em;
            font-weight: bold;
            border: none;
            border-radius: 12px;
            cursor: pointer;
            transition: all 0.3s;
            box-shadow: 0 4px 12px rgba(0,0,0,0.2);
        }
        
        .btn-primary {
            background: linear-gradient(135deg, #3f51b5 0%, #303f9f 100%);
            color: white;
        }
        
        .btn-primary:hover {
            transform: translateY(-2px);
            box-shadow: 0 6px 16px rgba(63, 81, 181, 0.4);
        }
        
        .btn-secondary {
            background: linear-gradient(135deg, #5c6bc0 0%, #3949ab 100%);
            color: white;
        }
        
        .btn-secondary:hover {
            transform: translateY(-2px);
            box-shadow: 0 6px 16px rgba(92, 107, 192, 0.4);
        }
        
        button:disabled {
            opacity: 0.6;
            cursor: not-allowed;
            transform: none !important;
        }
        
        .result-container {
            background: linear-gradient(135deg, rgba(197, 202, 233, 0.95) 0%, rgba(159, 168, 218, 0.95) 100%);
            border-radius: 20px;
            padding: 30px;
            margin-top: 25px;
            display: none;
            border: 3px solid #3f51b5;
        }
        
        .answer-box {
            background: white;
            padding: 20px;
            border-radius: 15px;
            margin-bottom: 15px;
            border-left: 4px solid #3f51b5;
        }
        
        .question-text {
            color: #303f9f;
            font-weight: bold;
            font-size: 1.1em;
            margin-bottom: 10px;
        }
        
        .answer-text {
            color: #1a237e;
            font-size: 1.3em;
            font-weight: bold;
            margin: 10px 0;
        }
        
        .confidence {
            color: #666;
            font-size: 0.95em;
        }
        
        .confidence-bar {
            width: 100%;
            height: 8px;
            background: #e8eaf6;
            border-radius: 4px;
            margin-top: 8px;
            overflow: hidden;
        }
        
        .confidence-fill {
            height: 100%;
            background: linear-gradient(90deg, #3f51b5 0%, #303f9f 100%);
            border-radius: 4px;
            transition: width 0.5s ease;
        }
        
        .loading {
            display: none;
            text-align: center;
            padding: 20px;
        }
        
        .spinner {
            border: 4px solid rgba(63, 81, 181, 0.1);
            border-left-color: #3f51b5;
            border-radius: 50%;
            width: 50px;
            height: 50px;
            animation: spin 1s linear infinite;
            margin: 0 auto 15px;
        }
        
        @keyframes spin {
            to { transform: rotate(360deg); }
        }
        
        .info-box {
            background: linear-gradient(135deg, #e8eaf615 0%, #c5cae915 100%);
            border-left: 4px solid #3f51b5;
            padding: 15px 20px;
            border-radius: 8px;
            margin-top: 20px;
        }
        
        .info-box p {
            color: #1a237e;
            line-height: 1.8;
            margin-bottom: 8px;
            font-size: 1.1em;
            font-weight: 500;
        }
        
        .info-box strong {
            color: #0d47a1;
            font-weight: 700;
        }
    </style>
</head>
<body>
    <div class="container">
        <h1>📄 文档理解</h1>
        <p class="subtitle">文档解析师帮你理解文档内容！</p>
        
        <div class="upload-section">
            <div class="upload-area" id="uploadArea">
                <div class="upload-icon">📄</div>
                <div>点击或拖拽文档到这里</div>
                <div style="color: #666; font-size: 0.9em; margin-top: 8px;">支持 PDF、Word、图片等格式</div>
                <input type="file" id="fileInput" accept="image/*,.pdf,.doc,.docx">
            </div>
            
            <div class="preview-container" id="previewContainer">
                <img id="previewImage" class="preview-image" alt="文档预览">
                
                <div class="question-area">
                    <label>❓ 向文档提问（支持中英文）：</label>
                    <input type="text" id="questionInput" placeholder="例如：总金额是多少？或 What is the total amount?">
                    <div class="quick-questions">
                        <button class="quick-btn" onclick="setQuestion('发票号是多少？')">发票号</button>
                        <button class="quick-btn" onclick="setQuestion('日期是什么？')">日期</button>
                        <button class="quick-btn" onclick="setQuestion('总金额是多少？')">总金额</button>
                        <button class="quick-btn" onclick="setQuestion('供应商是谁？')">供应商</button>
                        <button class="quick-btn" onclick="setQuestion('What is the name?')">姓名</button>
                        <button class="quick-btn" onclick="setQuestion('What is the phone number?')">电话</button>
                        <button class="quick-btn" onclick="setQuestion('What is the email?')">邮箱</button>
                        <button class="quick-btn" onclick="setQuestion('What is the invoice number?')">Invoice#</button>
                        <button class="quick-btn" onclick="setQuestion('What is the total amount?')">Total</button>
                    </div>
                </div>
                
                <div class="button-group">
                    <button class="btn-primary" id="askBtn" onclick="askQuestion()">🔍 提问</button>
                    <button class="btn-secondary" id="changeBtn" onclick="changeDocument()">🔄 更换文档</button>
                </div>
            </div>
        </div>
        
        <div class="loading" id="loading">
            <div class="spinner"></div>
            <div style="color: #303f9f; font-size: 1.1em; font-weight: 600;">AI 正在分析文档...</div>
        </div>
        
        <div class="result-container" id="resultContainer"></div>
        
        <div class="info-box">
            <p><strong>🤖 模型：</strong>impira/layoutlm-document-qa</p>
            <p><strong>💡 功能：</strong>OCR文字识别 + 布局分析 + 信息提取</p>
            <p><strong>✅ 推荐格式：</strong>图片（JPG/PNG/BMP）- 效果最好！</p>
            <p><strong>📋 其他格式：</strong>PDF（需转换）、Word（建议截图后上传）</p>
            <p><strong>🌏 语言支持：</strong>支持中文和英文提问（推荐使用英文提问效果更好）</p>
            <p><strong>💡 使用提示：</strong>Word文档请先截图保存为图片，然后上传图片文件</p>
            <p><strong>📸 图片要求：</strong></p>
            <p style="margin-left: 20px;">• 分辨率：建议至少 1000x1000 像素</p>
            <p style="margin-left: 20px;">• 清晰度：文字清晰可读，避免模糊</p>
            <p style="margin-left: 20px;">• 对比度：黑色文字 + 白色背景效果最好</p>
            <p style="margin-left: 20px;">• 布局：简单的表格/表单效果最好</p>
            <p><strong>❓ 提问技巧：</strong></p>
            <p style="margin-left: 20px;">• 使用英文提问准确度更高</p>
            <p style="margin-left: 20px;">• 问题要具体明确，如 "What is the invoice number?"</p>
            <p style="margin-left: 20px;">• 避免复杂的复合问题</p>
        </div>
    </div>

    <script>
        const fallingItems = ['📄', '📃', '📋', '📑', '📊', '📈', '🔍', '🔎', '✨', '⭐', '🌟', '💫', '📌', '🔖', '💼', '🏢'];
        
        function createFallingItem() {
            const item = document.createElement('div');
            item.className = 'falling-item';
            item.textContent = fallingItems[Math.floor(Math.random() * fallingItems.length)];
            item.style.left = Math.random() * 100 + '%';
            item.style.animationDuration = (Math.random() * 3 + 4) + 's';
            item.style.fontSize = (Math.random() * 16 + 20) + 'px';
            document.body.appendChild(item);
            
            setTimeout(() => item.remove(), 7000);
        }
        
        // 初始创建10个飘落元素
        for(let i = 0; i < 10; i++) {
            setTimeout(createFallingItem, i * 150);
        }
        
        setInterval(createFallingItem, 150);
        
        let selectedFile = null;
        
        const uploadArea = document.getElementById('uploadArea');
        const fileInput = document.getElementById('fileInput');
        const previewContainer = document.getElementById('previewContainer');
        const previewImage = document.getElementById('previewImage');
        
        uploadArea.addEventListener('click', () => fileInput.click());
        
        fileInput.addEventListener('change', (e) => {
            handleFile(e.target.files[0]);
        });
        
        uploadArea.addEventListener('dragover', (e) => {
            e.preventDefault();
            uploadArea.style.background = 'rgba(63, 81, 181, 0.15)';
        });
        
        uploadArea.addEventListener('dragleave', () => {
            uploadArea.style.background = 'rgba(63, 81, 181, 0.05)';
        });
        
        uploadArea.addEventListener('drop', (e) => {
            e.preventDefault();
            uploadArea.style.background = 'rgba(63, 81, 181, 0.05)';
            handleFile(e.dataTransfer.files[0]);
        });
        
        function handleFile(file) {
            if (!file) {
                alert('请选择文件！');
                return;
            }
            
            // 检查文件类型
            const validTypes = ['image/jpeg', 'image/png', 'image/bmp', 'image/gif', 
                              'application/pdf', 
                              'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                              'application/msword'];
            
            const validExtensions = ['.jpg', '.jpeg', '.png', '.bmp', '.gif', '.pdf', '.doc', '.docx'];
            const fileName = file.name.toLowerCase();
            const hasValidExtension = validExtensions.some(ext => fileName.endsWith(ext));
            
            if (!validTypes.includes(file.type) && !hasValidExtension) {
                alert('请选择支持的文件格式：图片(JPG/PNG/BMP/GIF)、PDF或Word文档！');
                return;
            }
            
            selectedFile = file;
            
            // 根据文件类型显示不同的预览
            if (file.type.startsWith('image/')) {
                const reader = new FileReader();
                reader.onload = (e) => {
                    previewImage.src = e.target.result;
                    previewImage.style.display = 'block';
                    uploadArea.style.display = 'none';
                    previewContainer.style.display = 'block';
                    document.getElementById('resultContainer').style.display = 'none';
                };
                reader.readAsDataURL(file);
            } else {
                // PDF或Word文档，显示文件信息
                previewImage.style.display = 'none';
                uploadArea.style.display = 'none';
                previewContainer.style.display = 'block';
                document.getElementById('resultContainer').style.display = 'none';
                
                // 在预览区域显示文件信息
                const fileInfo = document.createElement('div');
                fileInfo.style.cssText = 'background: white; padding: 20px; border-radius: 15px; text-align: center; margin-bottom: 20px; border: 2px solid #3f51b5;';
                fileInfo.innerHTML = `
                    <div style="font-size: 3em; margin-bottom: 10px;">${file.name.endsWith('.pdf') ? '📕' : '📘'}</div>
                    <div style="color: #303f9f; font-weight: bold; font-size: 1.2em;">${file.name}</div>
                    <div style="color: #666; margin-top: 5px;">大小: ${(file.size / 1024).toFixed(2)} KB</div>
                    <div style="color: #666; margin-top: 5px;">类型: ${file.name.endsWith('.pdf') ? 'PDF文档' : 'Word文档'}</div>
                `;
                
                // 插入到预览容器的开头
                const container = document.getElementById('previewContainer');
                const firstChild = container.firstChild;
                container.insertBefore(fileInfo, firstChild);
            }
        }
        
        function changeDocument() {
            uploadArea.style.display = 'block';
            previewContainer.style.display = 'none';
            document.getElementById('resultContainer').style.display = 'none';
            fileInput.value = '';
            selectedFile = null;
        }
        
        function setQuestion(question) {
            document.getElementById('questionInput').value = question;
        }
        
        async function askQuestion() {
            if (!selectedFile) return;
            
            const question = document.getElementById('questionInput').value.trim();
            if (!question) {
                alert('请输入问题！');
                return;
            }
            
            const formData = new FormData();
            formData.append('image', selectedFile);
            formData.append('question', question);
            
            const loading = document.getElementById('loading');
            const resultContainer = document.getElementById('resultContainer');
            const askBtn = document.getElementById('askBtn');
            
            loading.style.display = 'block';
            resultContainer.style.display = 'none';
            askBtn.disabled = true;
            
            try {
                const response = await fetch('/ask', {
                    method: 'POST',
                    body: formData
                });
                
                const data = await response.json();
                
                if (data.success) {
                    displayResult(data);
                } else {
                    alert('分析失败：' + data.error);
                }
            } catch (error) {
                alert('请求失败：' + error.message);
            } finally {
                loading.style.display = 'none';
                askBtn.disabled = false;
            }
        }
        
        function displayResult(data) {
            const container = document.getElementById('resultContainer');
            const confidence = (data.score * 100).toFixed(1);
            
            let html = '<h3 style="color: #303f9f; margin-bottom: 20px; text-align: center;">✨ 分析结果</h3>';
            
            html += '<div class="answer-box">';
            html += `<div class="question-text">❓ ${data.question}</div>`;
            html += `<div class="answer-text">💡 ${data.answer}</div>`;
            html += `<div class="confidence">置信度: ${confidence}%</div>`;
            html += '<div class="confidence-bar">';
            html += `<div class="confidence-fill" style="width: ${confidence}%"></div>`;
            html += '</div>';
            html += '</div>';
            
            container.innerHTML = html;
            container.style.display = 'block';
        }
        
        document.getElementById('questionInput').addEventListener('keydown', function(e) {
            if (e.key === 'Enter') {
                askQuestion();
            }
        });
    </script>
</body>
</html>
"""

@app.route('/')
def index():
    return render_template_string(HTML_TEMPLATE)

@app.route('/static/background')
def background():
    if os.path.exists(BACKGROUND_PATH):
        return send_file(BACKGROUND_PATH, mimetype='image/png')
    else:
        return '', 404

@app.route('/ask', methods=['POST'])
def ask():
    try:
        if 'image' not in request.files:
            return jsonify({'success': False, 'error': '没有上传文件'})
        
        file = request.files['image']
        question = request.form.get('question', '')
        
        if not question:
            return jsonify({'success': False, 'error': '请输入问题'})
        
        filename = file.filename.lower()
        
        # 处理不同类型的文件
        if filename.endswith('.pdf'):
            # 处理PDF文件
            if not PDF_SUPPORT:
                return jsonify({'success': False, 'error': 'PDF支持未启用，请安装 pdf2image 和 poppler'})
            
            # 保存临时PDF文件
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_pdf:
                file.save(tmp_pdf.name)
                pdf_path = tmp_pdf.name
            
            try:
                # 转换PDF第一页为图片
                images = convert_from_path(pdf_path, first_page=1, last_page=1)
                if not images:
                    return jsonify({'success': False, 'error': 'PDF转换失败'})
                
                image = images[0].convert('RGB')
            finally:
                os.unlink(pdf_path)
                
        elif filename.endswith(('.doc', '.docx')):
            # 处理Word文档 - 将每页转换为图片
            if not DOCX_READ_SUPPORT:
                return jsonify({
                    'success': False, 
                    'error': 'Word支持未启用，请安装 python-docx\n安装命令: pip install python-docx'
                })
            
            try:
                # 方法1：使用docx2pdf + pdf2image（如果都安装了）
                if DOCX_SUPPORT and PDF_SUPPORT:
                    # 保存临时Word文件
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_docx:
                        file.save(tmp_docx.name)
                        docx_path = tmp_docx.name
                    
                    try:
                        # 转换为PDF
                        pdf_path = docx_path.replace('.docx', '.pdf')
                        docx_to_pdf(docx_path, pdf_path)
                        
                        # 转换PDF第一页为图片
                        images = convert_from_path(pdf_path, first_page=1, last_page=1)
                        if images:
                            image = images[0].convert('RGB')
                        else:
                            raise Exception("PDF转换失败")
                    finally:
                        if os.path.exists(docx_path):
                            os.unlink(docx_path)
                        if os.path.exists(pdf_path):
                            os.unlink(pdf_path)
                
                # 方法2：使用python-docx读取文本，创建简单图片
                else:
                    from docx import Document
                    from PIL import Image, ImageDraw, ImageFont
                    
                    # 读取Word文档
                    doc = Document(file.stream)
                    
                    # 提取所有文本
                    full_text = []
                    for para in doc.paragraphs:
                        if para.text.strip():
                            full_text.append(para.text)
                    
                    # 提取表格内容
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                            if row_text:
                                full_text.append(row_text)
                    
                    if not full_text:
                        return jsonify({
                            'success': False,
                            'error': '⚠️ Word文档中没有找到文本内容'
                        })
                    
                    # 创建文本图片
                    text_content = '\n'.join(full_text[:50])  # 最多50行
                    
                    # 创建白色背景图片
                    img_width = 1200
                    img_height = max(800, len(full_text[:50]) * 30 + 100)
                    image = Image.new('RGB', (img_width, img_height), 'white')
                    draw = ImageDraw.Draw(image)
                    
                    # 使用默认字体
                    try:
                        font = ImageFont.truetype("arial.ttf", 20)
                    except:
                        font = ImageFont.load_default()
                    
                    # 绘制文本
                    y_position = 50
                    for line in full_text[:50]:
                        # 处理长行
                        if len(line) > 80:
                            line = line[:80] + '...'
                        draw.text((50, y_position), line, fill='black', font=font)
                        y_position += 30
                    
                    print(f"✅ Word文档已转换为图片（{len(full_text)}行文本）")
                    
            except Exception as e:
                print(f"Word文档处理失败: {e}")
                import traceback
                traceback.print_exc()
                return jsonify({
                    'success': False,
                    'error': f'⚠️ Word文档处理失败\n\n错误：{str(e)}\n\n💡 建议：\n1. 将Word文档另存为PDF\n2. 或截图保存为图片后上传'
                })
                
        else:
            # 处理图片文件
            image = Image.open(file.stream).convert('RGB')
        
        # 执行文档问答 - 完全绕过pytesseract，直接传图片给模型
        print("📄 准备进行文档问答...")
        
        # 策略：直接传图片给模型，让模型使用内置OCR（LayoutLM模型支持）
        # 不再尝试手动OCR，避免编码问题
        try:
            print("直接使用模型进行文档问答（模型内置OCR）...")
            # 直接传图片，不提供word_boxes，让模型自己处理
            # 注意：这里不能传word_boxes参数，否则模型会尝试调用pytesseract
            result = doc_qa(image=image, question=question)
                
        except (UnicodeDecodeError, ValueError) as e:
            print(f"模型内置OCR编码错误: {e}")
            import traceback
            traceback.print_exc()
            return jsonify({
                'success': False,
                'error': '⚠️ OCR识别失败（编码问题）\n\n这是Tesseract在Windows上的已知问题。\n\n💡 解决方案：\n1. 下载中文语言包 chi_sim.traineddata\n2. 放到：D:\\transformers训练\\transformers-main\\预训练模型下载处\n3. 或使用纯英文文档\n\n下载地址：\nhttps://github.com/tesseract-ocr/tessdata/raw/main/chi_sim.traineddata'
            })
        except Exception as e:
            print(f"文档问答失败: {e}")
            import traceback
            traceback.print_exc()
            return jsonify({
                'success': False,
                'error': f'⚠️ 文档分析失败\n\n错误信息：{str(e)}\n\n💡 建议：\n1. 使用JPG或PNG格式\n2. 确保图片包含清晰的文字\n3. 尝试使用英文文档'
            })
        
        # 检查结果
        if not result or len(result) == 0:
            return jsonify({
                'success': False,
                'error': '模型未能从图片中找到答案。请确保：\n1. 图片清晰，文字可读\n2. 问题与图片内容相关\n3. 尝试使用英文提问'
            })
        
        return jsonify({
            'success': True,
            'question': question,
            'answer': result[0]['answer'],
            'score': float(result[0]['score'])
        })
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': f'分析失败: {str(e)}'
        })

if __name__ == '__main__':
    import webbrowser
    import threading
    
    print("\n" + "=" * 70)
    print("📄 启动文档解析师...")
    print("=" * 70)
    print("\n📍 访问地址: http://localhost:8001")
    print("🔍 文档解析师在这里等你~\n")
    
    def open_browser():
        import time
        time.sleep(1)
        webbrowser.open('http://localhost:8001')
    
    threading.Thread(target=open_browser, daemon=True).start()
    
    app.run(host='0.0.0.0', port=8001, debug=False)
